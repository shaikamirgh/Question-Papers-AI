import pytesseract
from PIL import Image
import google.generativeai as palm
from docx import Document
from docx.enum.text import WD_BREAK
import docx2pdf

docx_file = "QnA_AI.docx"
pdf_file = "QnA_AI.pdf"
#"C:/Users/shaik/OneDrive/Desktop/Pyweek2023/extraprojects"
#files_path = "C:/Users/shaik/OneDrive/Desktop/Pyweek2023/extraprojects/"

def perform_ocr(image_path):
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return text
    except Exception as e:
        return str(e)


def get_answers(question_text):
    try:
        palm.configure(api_key='AIzaSyDYTiHCkpFbjNB28PKKgCkhi-kpchwv8GA')  # Replace with your API key

        palm_prompt = """
        You are a language model that provides correct and detailed answers to questions with examples in order. Put the question number then mention the question explicitly and then in the next line "Ans: " and Answer.
        Output format:
        <Question number>. <Question>
        <Ans:> <Answer>

        sample output:
        1. What is the name of the capital of India?
        Ans: New Delhi

        (two new lines space for next question)

        2. Explain the concept of inheritance in OOP?
        Ans: Inheritance is the process by which one class takes the properties of another class. The class that inherits the properties of another class is called the derived class or child class. The class whose properties are inherited is called the base class or parent class.
        """        

        response = palm.chat(messages=palm_prompt+question_text)

        # Last contains the model's response:
        answer = response.last
        return answer
    except Exception as e:
        return str(e)


def create_document(questions, answers):
    doc = Document()
    doc.add_heading('Question and Answers Document', 0)
    doc.add_paragraph(answers)
    docx_file = "QnA_AI.docx"
    pdf_file = "QnA_AI.pdf"
    doc.save(docx_file)





# Input: Provide the path to your question paper image
image_path = "extraprojects/cdqp2.jpeg"

# Step 1: Perform OCR
question_text = perform_ocr(image_path)
print(question_text)

print("--------------------------")
# Step 2: Get answers
print("Getting answers...")
answers = get_answers(question_text)
print(answers)
print("--------------------------")
# Step 3: Create a Word document
create_document(question_text, answers)
print("Document created: QnA_AI.docx")

print("converting to pdf")

docx2pdf.convert(docx_file, pdf_file)
#print full path
print("pdf created: ", pdf_file)

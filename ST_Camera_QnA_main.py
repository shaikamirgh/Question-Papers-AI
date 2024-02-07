import pytesseract
from PIL import Image
import google.generativeai as palm
from docx import Document
from docx.enum.text import WD_BREAK
# import docx2pdf
import streamlit as st
import time
import subprocess

docx_file = "QnA_AI.docx"
pdf_file = "QnA_AI.pdf"
folder_path = 'img_data/'
#files_path = "C:/Users/shaik/OneDrive/Desktop/Pyweek2023/extraprojects/"



# def docx_to_pdf(docx_file, pdf_file):
#     cmd = ['libreoffice', '--convert-to', 'pdf', docx_file, '--outdir', pdf_file]
#     subprocess.run(cmd)



def perform_ocr(image_path):
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return text
    except Exception as e:
        return str(e)


def get_answers(question_text):
    try:
        palm.configure(api_key='AIzaSyDYTiHCkpFbjNB28PKKgCkhi-kpchwv8GA')  # Replace with your API key

        palm_prompt = """
        You are a language model that provides correct and DETAILED answers to questions with examples in order. Put the question number then mention the question explicitly and then in the next line "Ans: " and Answer.
        Output format:
        <Question number>. <Question>
        <Ans:> <Answer>

        sample output:
        1. What is the difference between a constructor and a method??
        Ans: The main difference between a constructor and a method is that a constructor is called when an object is created, while a method is called after an object has been created.

        (two new lines space for next question)

        2. Explain the concept of inheritance in OOP?
        Ans: Inheritance is the process by which one class takes the properties of another class. The class that inherits the properties of another class is called the derived class or child class. The class whose properties are inherited is called the base class or parent class.
        """        

        response = palm.chat(messages=palm_prompt+question_text)

        # Last contains the model's response:
        answer = response.last
        return answer
    except Exception as e:
        return str(e)


def create_document(questions, answers):
    doc = Document()
    doc.add_heading('Question and Answers Document', 0)
    doc.add_paragraph(answers)
    docx_file = "QnA_AI.docx"
    pdf_file = "QnA_AI.pdf"
    doc.save(docx_file)
    #display this docx file in streamlit
    with open('QnA_AI.docx', 'rb') as f:
        st.download_button('Download Docx with Answers', f, file_name='QnA_AI.docx', key='docx_download_button')
    # Inject JavaScript to automatically click the download link
    # js_code = """
    # <script>
    # document.addEventListener("DOMContentLoaded", function() {
    #     const downloadLink = document.getElementById("docx_download_button").getElementsByTagName("a")[0];
    #     downloadLink.click();
    # });
    # </script>
    # """
    # st.markdown(js_code, unsafe_allow_html=True)

    return docx_file, pdf_file

def main():
    st.title("AI Assistant")
    st.write("Take a pic of a Question Paper (previous years obv) and upload here: ")
    img = st.file_uploader("Choose an image", type=["jpg", "jpeg", "png"])
    time.sleep(0.8)
    if img is not None:
        image = Image.open(img)
        st.image(image, caption='Uploaded Image.', use_column_width=True)
        print(type(image))
        image.save('img_data/capture.png')
        
        timestamp = int(time.time())
        file_extension = image.format.lower()
        file_path = f'{folder_path}{timestamp}.{file_extension}'
        image.save(file_path)
    else:
        image = Image.open('img_data/default1.jpg')
        st.image(image, caption='Default Image.', use_column_width=True)
        image.save('img_data/capture.png')
        
    # Step 1: Perform OCR
    question_text = perform_ocr('img_data/capture.png')
    print(question_text)
    st.subheader("Scroll down to Download the AI Generated QnA docx file.")
    
    #st.write("Scroll down to Download the AI Generated QnA docx file. (PDF version coming soon In Sha' Allah 😉)")
    st.write("\n\nThe questions detected are: \n\n", question_text)
    print("--------------------------")

    # Step 2: Get answers
    print("Getting answers...")
    st.subheader("Generating answers...")
    answers = get_answers(question_text)
    print(answers)
    st.write("\n\nThe Answers generated are: \n", answers)
    st.write("\nPDF version coming soon In Sha' Allah 😉\n")
    print("--------------------------")

    # Step 3: Create a Word document
    create_document(question_text, answers)
    print("Document created: QnA_AI.docx")
    print("converting to pdf")
    print("pdf created: ", pdf_file)

    # Step 4: Display additional images
    st.subheader("Additional Sample Papers: ")
    sample_img1="img_data/default1.jpg"
    sample_img2="img_data/cap.png"
    with open(sample_img1, "rb") as file1, open(sample_img2, "rb") as file2:
        image1 = Image.open(file1)
        image2 = Image.open(file2)

        col1, col2 = st.columns(2)
        with col1:
            st.image(image1, caption='Sample Image 1', use_column_width=True)
        with col2:
            st.image(image2, caption='Sample Image 2', use_column_width=True)

        # st.subheader("Scroll down to Download the AI Generated QnA docx file.")
  
if __name__ == '__main__':
    main()